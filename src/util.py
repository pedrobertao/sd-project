import docx
import re
import hashlib

def create_document(docId=None,doc_type=None, topic=None, author=None, 
    date=None, signature=None ,fields=None, admSign=None, text="Sem corpo deste documento"):

    print(admSign)
    if admSign!='None':
        admSign = "Documento validado pelo Autor"
    else:
        admSign = "Documento precisa ser valido pelo Autor"

    if doc_type == 'Requerimento' or doc_type =='Documento':
        document = docx.Document('./media/templates/doc_template1.docx')
        for paragraph in document.paragraphs:
            if(doc_type!=None):
                paragraph.text = re.sub(r'__TYPE__', doc_type, paragraph.text)
            if(author!=None):
                paragraph.text = re.sub(r'__AUTHOR__', author, paragraph.text)
            if(date!=None):
                paragraph.text = re.sub(r'__DATE__', date, paragraph.text)
            if(signature!=None):
                paragraph.text = re.sub(r'__SIGNATURE__', signature, paragraph.text)
            if(text!=None):
                paragraph.text = re.sub(r'__TEXT__', text, paragraph.text)
            if(fields!=None):   
                fieldsStr = fields.split(',')
                paragraph.text = re.sub(r'__SIGNBY__', fieldsStr[0],paragraph.text)
                paragraph.text = re.sub(r'__FIELD0__', fieldsStr[1],paragraph.text)
                paragraph.text = re.sub(r'__FIELD1__', fieldsStr[2],paragraph.text)
            paragraph.text = re.sub(r'__CARIMBO__', admSign,paragraph.text)

    if doc_type == 'Memorando' or doc_type == 'Atestado':
        document = docx.Document('./media/templates/doc_template2.docx')
        for paragraph in document.paragraphs:
            if(doc_type!=None):
                paragraph.text = re.sub(r'__TYPE__', doc_type, paragraph.text)
            if(author!=None):
                paragraph.text = re.sub(r'__AUTHOR__', author, paragraph.text)
            if(date!=None):
                paragraph.text = re.sub(r'__DATE__', date, paragraph.text)

    

    title = topic+'-'+docId
    title = re.sub(' ','_',title)
    #document.save('{}.docx'.format(title))
    document.save('./media/docs/{}.docx'.format(title))
    return './media/docs/{}.docx'.format(title)


def encrypt_string(hash_string):
    if not hash_string:
        return "Documento incompleto"
    sha_signature = \
        hashlib.sha256(hash_string.encode()).hexdigest()
    return sha_signature

def compare_signature(author, date, test_signature):
    valid_signature = encrypt_string(author + date)
    if(test_signature == valid_signature):
        return True
    else: False
